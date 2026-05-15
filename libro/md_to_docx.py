#!/usr/bin/env python3
"""Convertidor Markdown → DOCX para el manuscrito del libro WQuestions.
Maneja: títulos (# ## ###), párrafos, bloques de código (```), listas (- 1.),
negrita (**), itálica (*), tablas (| | |), enlaces [texto](url).
"""

import re
import sys
import glob
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE


# Caracteres matemáticos / técnicos que muchas fuentes corrientes (Calibri)
# no incluyen. Cuando aparecen, forzamos un fallback que sí los tiene.
MATH_CHARS = set("∪∩∈∉⊆⊂⊇⊃→←↔⇒⇐⇔↦∀∃∄∅∧∨¬≠≡≈≤≥≪≫±∓∇∂∫∑∏√∞⊕⊗⊥∥ℝℕℤℚℂℙ"
                 "αβγδεζηθικλμνξοπρστυφχψω"
                 "ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ"
                 "·×÷")
MATH_FONT = "Cambria Math"  # disponible por defecto en Word y con cobertura completa


def _add_text_with_math_fallback(para, text, *,
                                 bold=False, italic=False, underline=False,
                                 font_name=None, font_size=None, font_color=None):
    """Agrega `text` al párrafo, dividiendo en runs separados cuando aparecen
    caracteres matemáticos para que esos segmentos usen una fuente que los
    soporte. Conserva estilos (bold/italic/etc.) en todos los runs resultantes.
    """
    if not text:
        return
    # Si el texto ya iba en una fuente especial (code, etc.) y no contiene math,
    # no hace falta partir.
    if font_name and not any(c in MATH_CHARS for c in text):
        run = para.add_run(text)
        run.bold, run.italic, run.underline = bold, italic, underline
        run.font.name = font_name
        if font_size: run.font.size = font_size
        if font_color: run.font.color.rgb = font_color
        return

    # Partir el texto en segmentos consecutivos del mismo "tipo" (math vs no math)
    buf = []
    is_math = None
    def flush():
        nonlocal buf, is_math
        if not buf: return
        seg = "".join(buf)
        run = para.add_run(seg)
        run.bold, run.italic, run.underline = bold, italic, underline
        if is_math:
            run.font.name = MATH_FONT
        elif font_name:
            run.font.name = font_name
        if font_size: run.font.size = font_size
        if font_color: run.font.color.rgb = font_color
        buf = []

    for c in text:
        c_is_math = c in MATH_CHARS
        if is_math is None:
            is_math = c_is_math
        elif c_is_math != is_math:
            flush()
            is_math = c_is_math
        buf.append(c)
    flush()


def add_runs_with_inline_formatting(para, text):
    """Procesa **negrita**, *itálica*, `code`, enlaces, y agrega runs al
    párrafo. Los símbolos matemáticos (∪, ∩, →, ∈, etc.) se renderizan
    automáticamente en una fuente con cobertura completa.
    """
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            _add_text_with_math_fallback(para, text[pos:m.start()])
        token = m.group()
        if token.startswith('**'):
            _add_text_with_math_fallback(para, token[2:-2], bold=True)
        elif token.startswith('*'):
            _add_text_with_math_fallback(para, token[1:-1], italic=True)
        elif token.startswith('`'):
            _add_text_with_math_fallback(para, token[1:-1],
                                         font_name='Courier New',
                                         font_size=Pt(10))
        elif token.startswith('['):
            link_text = re.match(r'\[([^\]]+)\]', token).group(1)
            _add_text_with_math_fallback(para, link_text,
                                         underline=True,
                                         font_color=RGBColor(0x05, 0x63, 0xC1))
        pos = m.end()
    if pos < len(text):
        _add_text_with_math_fallback(para, text[pos:])


def convert_md_file(md_path, doc, is_first=True):
    """Procesa un .md y agrega su contenido a un Document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    if not is_first:
        doc.add_page_break()

    i = 0
    in_code = False
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Bloque de código
        if line.startswith('```'):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            # Partir la línea: segmentos sin math van en Courier New, los que
            # tienen math van en Cambria Math para asegurar el render.
            buf = []
            is_math = None
            def _flush_code(_p, _buf, _is_math):
                if not _buf: return
                r = _p.add_run("".join(_buf))
                r.font.name = MATH_FONT if _is_math else 'Courier New'
                r.font.size = Pt(9)
            for c in line:
                cm = c in MATH_CHARS
                if is_math is None:
                    is_math = cm
                elif cm != is_math:
                    _flush_code(para, buf, is_math)
                    buf = []
                    is_math = cm
                buf.append(c)
            _flush_code(para, buf, is_math)
            i += 1
            continue

        # Tabla
        if line.startswith('|') and '|' in line[1:]:
            cells = [c.strip() for c in line.strip('|').split('|')]
            # Saltar la línea separadora ---
            if all(set(c.replace('-', '').replace(':', '').strip()) <= set('') for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            in_table = True
            i += 1
            # Si la siguiente no es de tabla, escribir tabla
            if i >= len(lines) or not lines[i].rstrip('\n').startswith('|'):
                if table_rows:
                    n_cols = len(table_rows[0])
                    table = doc.add_table(rows=len(table_rows), cols=n_cols)
                    table.style = 'Light Grid Accent 1'
                    for r_idx, row in enumerate(table_rows):
                        for c_idx, cell_text in enumerate(row[:n_cols]):
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.paragraphs[0].text = ''
                            add_runs_with_inline_formatting(cell.paragraphs[0], cell_text)
                    table_rows = []
                    in_table = False
                    doc.add_paragraph()
            continue

        # Imagen markdown: ![alt](ruta)  o  ![alt](ruta "leyenda")
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)\s]+)(?:\s+"([^"]+)")?\)\s*$', line)
        if img_match:
            alt, src, caption = img_match.groups()
            # Resolver ruta relativa al archivo .md
            img_path = src if os.path.isabs(src) else os.path.normpath(
                os.path.join(os.path.dirname(md_path), src))
            if os.path.exists(img_path):
                para = doc.add_paragraph()
                para.alignment = 1  # centrado
                run = para.add_run()
                # Ancho máximo: 6 pulgadas (cabe en la página)
                run.add_picture(img_path, width=Inches(6))
                if caption or alt:
                    cap_para = doc.add_paragraph()
                    cap_para.alignment = 1
                    cap_run = cap_para.add_run(caption or alt)
                    cap_run.italic = True
                    cap_run.font.size = Pt(9)
            else:
                # Fallback: mostrar el alt como texto si la imagen no existe
                para = doc.add_paragraph()
                run = para.add_run(f"[imagen no encontrada: {src}]")
                run.italic = True
                run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
            i += 1
            continue

        # Títulos
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            h = doc.add_heading(line[5:], level=4)
        # Cita
        elif line.startswith('> '):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4)
            run = para.add_run(line[2:])
            run.italic = True
        # Lista con guión
        elif re.match(r'^\s*-\s', line):
            text = re.sub(r'^\s*-\s', '', line)
            para = doc.add_paragraph(style='List Bullet')
            add_runs_with_inline_formatting(para, text)
        # Lista numerada
        elif re.match(r'^\s*\d+\.\s', line):
            text = re.sub(r'^\s*\d+\.\s', '', line)
            para = doc.add_paragraph(style='List Number')
            add_runs_with_inline_formatting(para, text)
        # Línea horizontal
        elif line.strip() == '---':
            doc.add_paragraph('_' * 40)
        # Línea vacía
        elif not line.strip():
            pass
        # Párrafo normal
        else:
            para = doc.add_paragraph()
            add_runs_with_inline_formatting(para, line)

        i += 1


def main():
    libro_dir = os.path.dirname(os.path.abspath(__file__))
    manuscrito_dir = os.path.join(libro_dir, 'manuscrito')
    output_dir = os.path.join(libro_dir, 'docx')
    os.makedirs(output_dir, exist_ok=True)

    # Documentos a generar
    targets = []

    # Manuscrito completo
    md_files = sorted(glob.glob(os.path.join(manuscrito_dir, '*.md')))
    if md_files:
        doc = Document()
        # Página de título
        title = doc.add_heading('Las preguntas como coordenadas', level=0)
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(
            'Arquitectura universal para la información en la era de la IA'
        )
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(14)
        doc.add_paragraph()
        doc.add_page_break()

        for idx, md in enumerate(md_files):
            convert_md_file(md, doc, is_first=(idx == 0))

        # Anexo: referencias al final del manuscrito completo
        referencias_md = os.path.join(libro_dir, 'referencias.md')
        if os.path.exists(referencias_md):
            convert_md_file(referencias_md, doc, is_first=False)

        out = os.path.join(output_dir, 'manuscrito_completo.docx')
        doc.save(out)
        targets.append(out)

    # Propuesta editorial
    propuesta = os.path.join(libro_dir, 'propuesta_editorial.md')
    if os.path.exists(propuesta):
        doc = Document()
        convert_md_file(propuesta, doc, is_first=True)
        out = os.path.join(output_dir, 'propuesta_editorial.docx')
        doc.save(out)
        targets.append(out)

    # Esquema de capítulos
    esquema = os.path.join(libro_dir, 'esquema_capitulos.md')
    if os.path.exists(esquema):
        doc = Document()
        convert_md_file(esquema, doc, is_first=True)
        out = os.path.join(output_dir, 'esquema_capitulos.docx')
        doc.save(out)
        targets.append(out)

    # Referencias (documento independiente)
    referencias = os.path.join(libro_dir, 'referencias.md')
    if os.path.exists(referencias):
        doc = Document()
        convert_md_file(referencias, doc, is_first=True)
        out = os.path.join(output_dir, 'referencias.docx')
        doc.save(out)
        targets.append(out)

    # Capítulos individuales también, por si los quiere editar separados
    for md in md_files:
        doc = Document()
        convert_md_file(md, doc, is_first=True)
        nombre = os.path.basename(md).replace('.md', '.docx')
        out = os.path.join(output_dir, nombre)
        doc.save(out)
        targets.append(out)

    print(f"Generados {len(targets)} documentos en {output_dir}:")
    for t in targets:
        print(f"  ✓ {os.path.basename(t)}")


if __name__ == '__main__':
    main()
